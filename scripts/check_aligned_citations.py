from pathlib import Path
from zipfile import ZipFile

from docx import Document


DOCS = Path(r"D:\TA Skripsi\aplikasi\HeightMeasurement\output\documents")
checks = [
    ("LAPORAN_SKRIPSI_ROBUSTNESS_SITASI_SELARAS_FINAL.docx", [85, 86, 87, 88]),
    ("JURNAL_ROBUSTNESS_SITASI_SELARAS_FINAL.docx", [24, 25, 26, 27, 30]),
]

for filename, indexes in checks:
    path = DOCS / filename
    doc = Document(path)
    print(f"\n{filename}: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")
    for index in indexes:
        paragraph = doc.paragraphs[index]
        sdt_count = len(paragraph._p.xpath(".//w:sdt"))
        xml_text = "".join(paragraph._p.xpath(".//w:t/text()"))
        print(f"p{index}: sdt={sdt_count}; equal={paragraph.text == xml_text}; {paragraph.text}")
    with ZipFile(path) as archive:
        print("document.xml present:", "word/document.xml" in archive.namelist())
        document_xml = archive.read("word/document.xml").decode("utf-8")
    full_text = "\n".join(p.text for p in doc.paragraphs)
    expected_authors = (
        "Ulijaszek", "Graybeal", "Zhang", "Prasanna Pabba", "Dill", "Garrido-Jurado"
    )
    print("author occurrence counts:")
    for author in expected_authors:
        print(f"  {author}: visible={full_text.count(author)}, xml={document_xml.count(author)}")
