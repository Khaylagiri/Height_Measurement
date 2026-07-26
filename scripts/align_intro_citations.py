from copy import deepcopy
from pathlib import Path

from docx import Document


ROOT = Path(r"D:\TA Skripsi\aplikasi\HeightMeasurement")
DOCS = ROOT / "output" / "documents"


def replace_paragraph(paragraph, text):
    """Replace all paragraph content, including citation content controls."""
    first_rpr = None
    for run in paragraph.runs:
        if run._r.rPr is not None:
            first_rpr = deepcopy(run._r.rPr)
            break

    ppr = paragraph._p.pPr
    for child in list(paragraph._p):
        if child is not ppr:
            paragraph._p.remove(child)

    run = paragraph.add_run(text)
    if first_rpr is not None:
        if run._r.rPr is not None:
            run._r.remove(run._r.rPr)
        run._r.insert(0, first_rpr)


def insert_reference_before(paragraph, text):
    new_paragraph = paragraph.insert_paragraph_before(text, style=paragraph.style)
    if paragraph._p.pPr is not None:
        current_ppr = new_paragraph._p.pPr
        if current_ppr is not None:
            new_paragraph._p.remove(current_ppr)
        new_paragraph._p.insert(0, deepcopy(paragraph._p.pPr))
    if paragraph.runs and paragraph.runs[0]._r.rPr is not None and new_paragraph.runs:
        new_run = new_paragraph.runs[0]
        if new_run._r.rPr is not None:
            new_run._r.remove(new_run._r.rPr)
        new_run._r.insert(0, deepcopy(paragraph.runs[0]._r.rPr))
    return new_paragraph


def append_reference(doc, text):
    template = doc.paragraphs[-1]
    paragraph = doc.add_paragraph(text, style=template.style)
    if template._p.pPr is not None:
        current_ppr = paragraph._p.pPr
        if current_ppr is not None:
            paragraph._p.remove(current_ppr)
        paragraph._p.insert(0, deepcopy(template._p.pPr))
    if template.runs and template.runs[0]._r.rPr is not None and paragraph.runs:
        run = paragraph.runs[0]
        if run._r.rPr is not None:
            run._r.remove(run._r.rPr)
        run._r.insert(0, deepcopy(template.runs[0]._r.rPr))
    return paragraph


def edit_report():
    source = DOCS / "LAPORAN_SKRIPSI_ROBUSTNESS_SUDUT_30.docx"
    target = DOCS / "LAPORAN_SKRIPSI_ROBUSTNESS_SITASI_SELARAS_FINAL.docx"
    doc = Document(source)

    replacements = {
        85: (
            "Pengukuran tinggi badan merupakan salah satu parameter antropometri yang penting "
            "dalam bidang kesehatan dan olahraga. Dalam bidang kesehatan, tinggi badan digunakan "
            "untuk menilai status gizi, memantau pertumbuhan, dan menghitung indeks massa tubuh "
            "(Aryani et al., 2023). Dalam bidang olahraga, data antropometri mendukung evaluasi "
            "dan pembinaan atlet sesuai karakteristik cabang olahraga (Azizah dan Nurrochmah, "
            "2024). Oleh karena itu, pengukuran tinggi badan yang akurat dan efisien menjadi "
            "kebutuhan penting."
        ),
        86: (
            "Metode pengukuran tinggi badan secara konvensional umumnya menggunakan stadiometer "
            "atau microtoise. Meskipun cukup akurat, hasil pengukuran dapat dipengaruhi oleh "
            "prosedur, posisi subjek, serta variasi antar-pengukur (Ulijaszek dan Kerr, 1999). "
            "Selain itu, keterbatasan alat dan tenaga terlatih dapat mengurangi efisiensi "
            "pengukuran pada kondisi tertentu (Roberto Fernandes da Costa et al., 2023)."
        ),
        87: (
            "Perkembangan teknologi computer vision pada perangkat mobile membuka peluang untuk "
            "melakukan pengukuran antropometri secara lebih praktis dan mandiri (Graybeal et al., "
            "2023). Dalam pendekatan ini, perspective correction digunakan untuk mengurangi "
            "distorsi akibat perbedaan sudut pandang kamera (Zhang et al., 2024), sedangkan pose "
            "estimation seperti MediaPipe Pose membantu mendeteksi landmark tubuh secara otomatis "
            "(Prasanna Pabba et al., 2025). Kinerja pose estimation juga dapat dipengaruhi oleh "
            "sudut kamera (Dill et al., 2023)."
        ),
        88: (
            "Namun, penerapan pengukuran tinggi badan berbasis citra pada perangkat mobile masih "
            "menghadapi tantangan. Variasi pengguna, lingkungan, dan perangkat dapat memengaruhi "
            "hasil pengukuran (Ma et al., 2024). Beberapa metode juga masih memerlukan perangkat "
            "tambahan atau prosedur kalibrasi yang kompleks (Xiao et al., 2025). Oleh karena itu, "
            "integrasi perspective correction dan pose estimation masih perlu dievaluasi secara "
            "terkontrol pada perangkat mobile."
        ),
    }
    for index, text in replacements.items():
        replace_paragraph(doc.paragraphs[index], text)

    placeholder = next(
        paragraph for paragraph in doc.paragraphs
        if paragraph.text.startswith("[LENGKAPI DAFTAR PUSTAKA")
    )
    insert_reference_before(
        placeholder,
        "Ulijaszek, S. J., dan Kerr, D. A. (1999). Anthropometric Measurement Error "
        "and the Assessment of Nutritional Status. British Journal of Nutrition, "
        "82(3), 165–177. https://doi.org/10.1017/S0007114599001348"
    )

    doc.save(target)
    return target


def edit_journal():
    source = DOCS / "JURNAL_ROBUSTNESS_SUDUT_30.docx"
    target = DOCS / "JURNAL_ROBUSTNESS_SITASI_SELARAS_FINAL.docx"
    doc = Document(source)

    replacements = {
        24: (
            "Tinggi badan merupakan salah satu parameter antropometri dasar yang digunakan dalam "
            "penilaian status gizi, pemantauan pertumbuhan, dan evaluasi kesehatan. Ketepatan "
            "pengukuran dipengaruhi oleh prosedur, posisi subjek, dan variasi antar-pengamat "
            "(Ulijaszek dan Kerr, 1999)."
        ),
        25: (
            "Perkembangan kamera smartphone dan computer vision membuka peluang pengukuran "
            "antropometri yang lebih praktis dan mandiri (Graybeal et al., 2023). Namun, ketepatan "
            "hasil pengukuran berbasis citra tetap perlu diuji terhadap variasi kondisi "
            "pengambilan gambar."
        ),
        26: (
            "Pengukuran tinggi badan dari citra memerlukan skala referensi untuk mengonversi "
            "ukuran piksel menjadi satuan panjang. Ketika kamera tidak sejajar dengan subjek, "
            "perbedaan sudut pandang dapat menimbulkan distorsi perspektif (Zhang et al., 2024). "
            "Selain itu, ketepatan pose estimation dapat dipengaruhi oleh sudut kamera dan "
            "konfigurasi gerak yang dianalisis (Dill et al., 2023)."
        ),
        27: (
            "ArUco marker dapat digunakan sebagai referensi skala karena memiliki identitas unik "
            "dan empat titik sudut yang dapat dideteksi pada citra. Informasi geometris marker "
            "tersebut mendukung proses kalibrasi dan perspective correction (Garrido-Jurado et "
            "al., 2014)."
        ),
        30: (
            "Pada penelitian ini, robustness didefinisikan sebagai kemampuan sistem mempertahankan "
            "akurasi dan keberhasilan pengukuran ketika sudut horizontal kamera berubah pada "
            "−30°, −15°, 0°, +15°, dan +30°, sementara jarak kamera, tinggi tripod, zoom, "
            "pencahayaan, latar belakang, dan postur subjek dikendalikan. Robustness dievaluasi "
            "melalui perbandingan mean absolute error (MAE), perubahan MAE (ΔMAE), dan success "
            "rate antara pengukuran tanpa dan dengan perspective correction."
        ),
    }
    for index, text in replacements.items():
        replace_paragraph(doc.paragraphs[index], text)

    append_reference(
        doc,
        "Zhang, Z., Zhou, J., Li, X., Xu, C., Hu, X., dan Wang, L. (2024). Correction "
        "Method for Perspective Distortions of Pipeline Images. Electronics, 13(15), "
        "2898. https://doi.org/10.3390/electronics13152898"
    )

    doc.save(target)
    return target


if __name__ == "__main__":
    for output in (edit_report(), edit_journal()):
        print(output)
